import os
import tempfile
from werkzeug.utils import secure_filename
from docx import Document
from PyPDF2 import PdfReader
import subprocess

ALLOWED_DOCUMENT_EXTENSIONS = {'pdf', 'docx', 'txt', 'pptx'}

def allowed_document_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_DOCUMENT_EXTENSIONS

def convert_document(file, input_format, output_format):
    filename = secure_filename(file.filename)
    input_path = os.path.join(tempfile.gettempdir(), filename)
    file.save(input_path)

    name_only = os.path.splitext(filename)[0]
    output_filename = f"{name_only}.{output_format}"
    output_path = os.path.join(tempfile.gettempdir(), output_filename)

    try:
        # DOCX -> PDF using LibreOffice on Linux
        if input_format == "docx" and output_format == "pdf":
            command = [
                "soffice", "--headless", "--convert-to", "pdf",
                "--outdir", tempfile.gettempdir(), input_path
            ]
            subprocess.run(command, check=True, capture_output=True, text=True)
            
            converted_filename = f"{name_only}.pdf"
            converted_path = os.path.join(tempfile.gettempdir(), converted_filename)
            return converted_path, None

        # DOCX -> TXT
        if input_format == "docx" and output_format == "txt":
            doc = Document(input_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            return output_path, None

        # TXT -> DOCX
        elif input_format == "txt" and output_format == "docx":
            with open(input_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()
            doc = Document()
            for line in lines:
                doc.add_paragraph(line.strip())
            doc.save(output_path)
            return output_path, None
        
        # PDF -> TXT
        elif input_format == "pdf" and output_format == "txt":
            reader = PdfReader(input_path)
            text = ''
            for page in reader.pages:
                text += page.extract_text() + '\n'
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            return output_path, None
        
        else:
            return None, "Unsupported conversion. This feature is being updated."

    except Exception as e:
        return None, str(e)

    finally:
        if os.path.exists(input_path):
            os.remove(input_path)